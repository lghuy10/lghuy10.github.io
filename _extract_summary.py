import json, os, re, sys
from pathlib import Path

DOCX_PATH = Path(__file__).parent / "SummaryFes.docx"
OUT_MD = Path(__file__).parent / "temp_fes_extract" / "_SummaryFes_extracted.md"
OUT_JSON = Path(__file__).parent / "temp_fes_extract" / "_SummaryFes_structured.json"

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx chua cai. Cai bang: py -3 -m pip install python-docx")
    # Thử dùng zipfile fallback (đọc xml thô)
    import zipfile
    import xml.etree.ElementTree as ET
    ns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    print("Dùng fallback zipfile+xml docx (không cần python-docx)...")
    text_lines = []
    with zipfile.ZipFile(str(DOCX_PATH), 'r') as z:
        with z.open('word/document.xml') as f:
            tree = ET.parse(f)
            root = tree.getroot()
            for p in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                texts = []
                for t in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    texts.append(t.text or '')
                line = ''.join(texts).strip()
                if line:
                    text_lines.append(line)
    OUT_MD.parent.mkdir(exist_ok=True)
    with open(str(OUT_MD), 'w', encoding='utf-8') as f:
        for ln in text_lines:
            f.write("- " + ln + "\n")
    print(f"[FALLBACK OK] {len(text_lines)} lines -> {OUT_MD}")
    sys.exit(0)

doc = Document(str(DOCX_PATH))
total_images = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        total_images += 1

blocks = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name if para.style else ""
    hlvl = None
    m = re.match(r"Heading\s*(\d+)", style, re.IGNORECASE)
    if m:
        hlvl = int(m.group(1))
    else:
        if re.match(r"^[0-9]+[\.\)]", text):
            hlvl = 2
        elif re.match(r"^[•\-]", text):
            hlvl = 0
    blocks.append({
        "idx": i,
        "type": "heading" if hlvl else "para",
        "level": hlvl or 0,
        "text": text,
        "style": style,
    })

tables_out = []
for ti, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        rows.append(cells)
    tables_out.append({"rows": rows})

OUT_MD.parent.mkdir(exist_ok=True)
with open(str(OUT_MD), "w", encoding="utf-8") as f:
    f.write(f"# SummaryFes.docx\n\n")
    f.write(f"**Inline images:** {total_images}\n\n")
    f.write("## Blocks\n\n")
    for b in blocks:
        if b["type"] == "heading":
            f.write("\n" + ("#" * (b["level"] + 1)) + " " + b["text"] + "\n\n")
        else:
            f.write("- " + b["text"] + "\n\n")
    if tables_out:
        f.write("\n## Tables\n\n")
        for ti, t in enumerate(tables_out):
            f.write(f"### Table {ti+1}\n\n")
            for row in t["rows"]:
                f.write("| " + " | ".join(row) + " |\n")
            f.write("\n")

with open(str(OUT_JSON), "w", encoding="utf-8") as f:
    json.dump({"blocks": blocks, "tables": tables_out, "n_images": total_images}, f, ensure_ascii=False, indent=2)

print(f"[OK] python-docx read: blocks={len(blocks)}, tables={len(tables_out)}, imgs={total_images}")
print(f"     -> MD:   {OUT_MD}")
print(f"     -> JSON: {OUT_JSON}")
