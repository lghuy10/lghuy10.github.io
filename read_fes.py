"""Script trích xuất nội dung Word (docx) từ thư mục Fes và xuất ra markdown JSON."""
import json
import os
import re
from docx import Document
from docx.document import Document as Doc
from docx.oxml.ns import qn

FES_DIR = os.path.join(os.path.dirname(__file__), "Fes")
OUT_DIR = os.path.join(os.path.dirname(__file__), "temp_fes_extract")
os.makedirs(OUT_DIR, exist_ok=True)

FILES_TO_READ = [
    "Lễ hội Nghinh Ông Cần Giờ.docx",
    "Lễ hội Nguyên Tiêu.docx",
    "Lễ hội Thần Rừng.docx",
    "Lễ hội chùa Bà Thiên Hậu.docx",
    "Lễ hội Cúng Thần Lúa của Người Chơ - ro.docx",
    "Lễ hội Mùa Trái Chín.docx",
    "Lễ hội Khai hạ - Cầu an tại Lăng Tả quân Lê Văn Duyệt.docx",
    "Lễ Hội Dinh Cô Long Hải.docx",
    "Lễ Hội Nghinh Ông Tam Thắng.docx",
    "Lễ Hội Trần Hưng Đạo.docx",
]


def run_prop_text(run):
    return run.text or ""


def paragraph_styles_text(p):
    return p.style.name if p.style else ""


def extract_docx(path):
    doc = Document(path)
    blocks = []
    # Đếm số hình ảnh inline
    total_images = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            total_images += 1

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = paragraph_styles_text(para)
        heading_lvl = None
        m = re.match(r"Heading\s*(\d+)", style, re.IGNORECASE)
        if m:
            heading_lvl = int(m.group(1))
        else:
            # Thử dự đoán theo dấu đầu dòng
            if re.match(r"^[0-9]+[\.\)]", text):
                heading_lvl = 2
            elif re.match(r"^[•\-]", text):
                heading_lvl = 0
        blocks.append({
            "idx": i,
            "type": "heading" if heading_lvl else "para",
            "level": heading_lvl or 0,
            "text": text,
            "style": style,
        })

    # Lấy tất cả table cells (nếu có nội dung dạng bảng chứa các trường)
    tables_out = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)
        tables_out.append({"rows": rows})

    return {
        "file": os.path.basename(path),
        "total_images_detected": total_images,
        "blocks": blocks,
        "tables": tables_out,
    }


def main():
    results = []
    for name in FILES_TO_READ:
        path = os.path.join(FES_DIR, name)
        if not os.path.exists(path):
            results.append({"file": name, "missing": True})
            print(f"[SKIP] Không tìm thấy file: {name}")
            continue
        data = extract_docx(path)
        # Ghi file riêng để xem chi tiết
        out_md = os.path.join(OUT_DIR, name.replace(".docx", ".md"))
        with open(out_md, "w", encoding="utf-8") as f:
            f.write(f"# {name}\n\n")
            f.write(f"**Hình ảnh (số lượng inline relation):** {data['total_images_detected']}\n\n")
            f.write("## Nội dung blocks\n\n")
            for b in data["blocks"]:
                if b["type"] == "heading":
                    f.write(f"\n{'#' * (b['level'] + 1)} {b['text']}\n\n")
                else:
                    f.write(f"- {b['text']}\n\n")
            f.write("\n## Bảng (nếu có)\n\n")
            for ti, t in enumerate(data["tables"]):
                f.write(f"### Table {ti+1}\n\n")
                for row in t["rows"]:
                    f.write("| " + " | ".join(row) + " |\n")
                f.write("\n")
        results.append(data)
        print(f"[OK] Đã đọc: {name} -> blocks={len(data['blocks'])}, images={data['total_images_detected']}, tables={len(data['tables'])}")

    # Ghi JSON tổng hợp
    with open(os.path.join(OUT_DIR, "all_extracts.json"), "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"\n✅ Đã xuất ra thư mục: {OUT_DIR}")


if __name__ == "__main__":
    main()
